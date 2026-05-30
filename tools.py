# Copyright 2026 José Milán Carrasco
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.

"""
Herramientas para el chatbot
"""
import io
import os
import re
import uuid
from pathlib import Path
import httpx
from bs4 import BeautifulSoup
import PyPDF2
from docx import Document
from docx.shared import Pt
import yt_dlp
import asyncio
import logging
import urllib.parse
import json
import subprocess
import sys

import sqlite3

DB_FILE = 'user_profile.db'

def _init_db():
    try:
        with sqlite3.connect(DB_FILE) as conn:
            c = conn.cursor()
            c.execute('''
                CREATE TABLE IF NOT EXISTS interests (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    topic TEXT UNIQUE NOT NULL,
                    added_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            c.execute('''
                CREATE TABLE IF NOT EXISTS successful_scripts (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    intent_key TEXT UNIQUE NOT NULL,
                    code TEXT NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            conn.commit()
    except Exception as e:
        logging.error(f"Error inicializando BD: {e}")

# Initialize DB on load
_init_db()

def get_user_memory() -> str:
    """Lee el perfil del usuario de la base de datos."""
    try:
        with sqlite3.connect(DB_FILE) as conn:
            c = conn.cursor()
            c.execute('SELECT topic FROM interests ORDER BY added_at ASC')
            rows = c.fetchall()
            interests = [row[0] for row in rows]
            return ", ".join(interests)
    except Exception as e:
        logging.error(f"Error leyendo memoria DB: {e}")
        return ""

def update_user_memory(new_interests: list):
    """Añade nuevos intereses a la base de datos."""
    if not new_interests:
        return
    
    added = False
    try:
        with sqlite3.connect(DB_FILE) as conn:
            c = conn.cursor()
            for interest in new_interests:
                clean_interest = interest.strip().title()
                if clean_interest:
                    try:
                        c.execute('INSERT INTO interests (topic) VALUES (?)', (clean_interest,))
                        added = True
                    except sqlite3.IntegrityError:
                        # Ya existe el interés
                        pass
            if added:
                conn.commit()
                logging.info(f"Memoria actualizada en DB: {new_interests}")
    except Exception as e:
        logging.error(f"Error guardando memoria DB: {e}")

def clear_user_memory() -> bool:
    """Borra el perfil de usuario (memoria a largo plazo)."""
    try:
        with sqlite3.connect(DB_FILE) as conn:
            c = conn.cursor()
            c.execute('DELETE FROM interests')
            conn.commit()
        logging.info("Memoria a largo plazo borrada (DB).")
        return True
    except Exception as e:
        logging.error(f"Error borrando memoria DB: {e}")
        return False

def find_saved_script(prompt: str) -> str:
    """Busca si hay un script guardado que coincida con las palabras clave del prompt."""
    try:
        prompt_clean = prompt.lower().strip()
        intents = {
            "limpiar_temporales": ["limpia", "temporales", "borra", "temp", "archivo", "basura", "limpieza"],
            "recursos_sistema": ["recursos", "sistema", "cpu", "ram", "memoria", "procesador", "disco", "pantalla", "tarjeta", "grafica"],
        }
        matched_intent = None
        for intent, keywords in intents.items():
            matches = sum(1 for kw in keywords if kw in prompt_clean)
            if matches >= 2 or (intent == "limpiar_temporales" and any(x in prompt_clean for x in ["temporal", "temporales"])):
                matched_intent = intent
                break
        if not matched_intent:
            return None
        with sqlite3.connect(DB_FILE) as conn:
            c = conn.cursor()
            c.execute('SELECT code FROM successful_scripts WHERE intent_key = ?', (matched_intent,))
            row = c.fetchone()
            if row:
                return row[0]
    except Exception as e:
        logging.error(f"Error buscando script guardado: {e}")
    return None

def save_successful_script(prompt: str, code: str):
    """Guarda un script que se ejecutó correctamente si coincide con un intent conocido."""
    try:
        prompt_clean = prompt.lower().strip()
        intents = {
            "limpiar_temporales": ["limpia", "temporales", "borra", "temp", "archivo", "basura", "limpieza"],
            "recursos_sistema": ["recursos", "sistema", "cpu", "ram", "memoria", "procesador", "disco", "pantalla", "tarjeta", "grafica"],
        }
        matched_intent = None
        for intent, keywords in intents.items():
            matches = sum(1 for kw in keywords if kw in prompt_clean)
            if matches >= 2 or (intent == "limpiar_temporales" and any(x in prompt_clean for x in ["temporal", "temporales"])):
                matched_intent = intent
                break
        if matched_intent:
            with sqlite3.connect(DB_FILE) as conn:
                c = conn.cursor()
                c.execute('''
                    INSERT INTO successful_scripts (intent_key, code)
                    VALUES (?, ?)
                    ON CONFLICT(intent_key) DO UPDATE SET code = excluded.code
                ''', (matched_intent, code))
                conn.commit()
                logging.info(f"Script guardado con éxito para el intent: {matched_intent}")
    except Exception as e:
        logging.error(f"Error guardando script exitoso: {e}")

def extract_url_content(url: str) -> str:
    """Extrae el contenido principal de una URL de forma limpia."""
    try:
        # Asegurar que tiene esquema
        if not url.startswith('http'):
            url = 'https://' + url

        with httpx.Client(timeout=30.0, follow_redirects=True) as client:
            headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
            }
            response = client.get(url, headers=headers)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Eliminar ruido (scripts, estilos, nav, etc)
            for tag in soup(['script', 'style', 'nav', 'footer', 'header', 'aside', 'form']):
                tag.decompose()
            
            # Extraer texto con separador para mantener estructura básica
            text = soup.get_text(separator='\n')
            
            # Limpiar espacios y líneas vacías
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            clean_text = "\n".join(lines)
            
            # Limitar tamaño para no saturar el contexto de Ollama
            return clean_text[:8000] 
    except Exception as e:
        return f"[Error al intentar acceder a la web {url}: {str(e)}]"

async def async_extract_url_content(url: str) -> str:
    """Versión asíncrona de la extracción de URL."""
    try:
        if not url.startswith('http'):
            url = 'https://' + url

        async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
            headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
            }
            response = await client.get(url, headers=headers)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            for tag in soup(['script', 'style', 'nav', 'footer', 'header', 'aside', 'form']):
                tag.decompose()
            
            text = soup.get_text(separator='\n')
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            clean_text = "\n".join(lines)
            return clean_text[:8000] 
    except Exception as e:
        return f"[Error al intentar acceder a la web {url}: {str(e)}]"

async def async_search_web(query: str, max_results: int = 5) -> str:
    """Busca en internet usando DuckDuckGo (HTML) y devuelve un resumen de los resultados."""
    try:
        url = 'https://html.duckduckgo.com/html/'
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        }
        data = {'q': query}
        
        async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
            response = await client.post(url, headers=headers, data=data)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            results = []
            
            for result in soup.find_all('div', class_='result'):
                if len(results) >= max_results:
                    break
                
                title_a = result.find('a', class_='result__url')
                snippet_a = result.find('a', class_='result__snippet')
                
                if title_a and snippet_a:
                    title = title_a.text.strip()
                    href = title_a.get('href', '')
                    if 'uddg' in href:
                        parsed = urllib.parse.urlparse(href)
                        qs = urllib.parse.parse_qs(parsed.query)
                        if 'uddg' in qs:
                            href = qs['uddg'][0]
                    
                    snippet = snippet_a.text.strip()
                    results.append({"title": title, "href": href, "snippet": snippet})
            
            if not results:
                return "No se encontraron resultados para la búsqueda."
                
            output = "RESULTADOS DE BÚSQUEDA:\n"
            for r in results:
                output += f"Título: {r['title']}\nEnlace: {r['href']}\nResumen: {r['snippet']}\n\n"
            
            # Extraer el contenido real de los 2 primeros enlaces
            output += "--- CONTENIDO DETALLADO DE LOS PRIMEROS RESULTADOS ---\n"
            for r in results[:2]:
                try:
                    content = await async_extract_url_content(r['href'])
                    # Limitar a 3000 caracteres por enlace para no saturar
                    output += f"\n--- WEB: {r['title']} ---\n{content[:3000]}\n"
                except Exception:
                    pass
                
            return output
    except Exception as e:
        return f"[Error al buscar en internet: {str(e)}]"

def generate_word_file(markdown_text: str, filename: str = None) -> str:
    """Genera un archivo Word (.docx) a partir de texto Markdown mejorado."""
    if not os.path.exists('downloads'):
        os.makedirs('downloads')

    # Limpiar nombre de archivo sugerido
    if filename:
        filename = re.sub(r'[^\w\s.-]', '', filename).strip().replace(' ', '_')
        if not filename.endswith('.docx'):
            filename += '.docx'
    else:
        filename = f"doc_{uuid.uuid4().hex[:6]}.docx"
        
    filepath = os.path.join('downloads', filename)
    
    doc = Document()
    
    # Limpiar LaTeX y caracteres raros comunes de modelos AI
    clean_text = markdown_text
    clean_text = re.sub(r'\$\\text\{(.+?)\}\$', r'\1', clean_text)
    clean_text = re.sub(r'\$(.+?)\$', r'\1', clean_text)
    clean_text = re.sub(r'\\_', '_', clean_text)
    clean_text = re.sub(r'\\text\{(.+?)\}', r'\1', clean_text)
    
    lines = clean_text.splitlines()
    
    in_table = False
    table_rows = []

    def _render_table(doc_obj, rows_data):
        if not rows_data: return
        parsed_rows = []
        for r in rows_data:
            # Skip separator rows (e.g. |---|---|)
            if re.match(r'^\|[-\s:|]+\|$', r):
                continue
            cells = [c.strip() for c in r.split('|')]
            if cells and cells[0] == '': cells.pop(0)
            if cells and cells[-1] == '': cells.pop()
            parsed_rows.append(cells)
            
        if not parsed_rows: return
        
        max_cols = max(len(r) for r in parsed_rows)
        if max_cols == 0: return
        
        try:
            table = doc_obj.add_table(rows=len(parsed_rows), cols=max_cols)
            table.style = 'Table Grid'
        except Exception:
            table = doc_obj.add_table(rows=len(parsed_rows), cols=max_cols)
            
        for row_idx, row_data in enumerate(parsed_rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < max_cols:
                    cell = table.cell(row_idx, col_idx)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    parts = re.split(r'(\*\*.*?\*\*)', cell_text)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            p.add_run(part[2:-2]).bold = True
                        else:
                            p.add_run(part)

    for line in lines:
        line_strip = line.strip()
        
        # Detectar tablas markdown
        if line_strip.startswith('|') and line_strip.endswith('|'):
            in_table = True
            table_rows.append(line_strip)
            continue
        elif in_table:
            _render_table(doc, table_rows)
            in_table = False
            table_rows = []
            
        if not line_strip:
            doc.add_paragraph()
            continue
            
        # Detección de Cabeceras (con o sin espacio)
        if line_strip.startswith('# '):
            doc.add_heading(line_strip.lstrip('# ').strip(), level=1)
        elif line_strip.startswith('## '):
            doc.add_heading(line_strip.lstrip('# ').strip(), level=2)
        elif line_strip.startswith('### '):
            doc.add_heading(line_strip.lstrip('# ').strip(), level=3)
        elif line_strip.startswith('#### '):
            doc.add_heading(line_strip.lstrip('# ').strip(), level=4)
        
        # Listas
        elif line_strip.startswith('- ') or line_strip.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            # Limpiar negritas dentro de la lista
            text = line_strip[2:].replace('**', '').replace('__', '')
            p.add_run(text)
            
        elif re.match(r'^\d+\. ', line_strip):
            content = re.sub(r'^\d+\. ', '', line_strip)
            p = doc.add_paragraph(style='List Number')
            # Limpiar negritas
            text = content.replace('**', '').replace('__', '')
            p.add_run(text)
            
        else:
            # Párrafo normal con soporte básico de negritas
            p = doc.add_paragraph()
            # Dividir por negritas **texto**
            parts = re.split(r'(\*\*.*?\*\*)', line_strip)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
                    
    if in_table:
        _render_table(doc, table_rows)
            
    doc.save(filepath)
    return filename

def extract_text_from_file(file_bytes: bytes, file_type: str, file_name: str) -> str:
    """Extrae texto de PDF, DOCX o TXT."""
    try:
        if file_type == 'txt':
            return file_bytes.decode('utf-8', errors='replace')

        elif file_type == 'pdf':
            try:
                reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                pages = [page.extract_text() or '' for page in reader.pages]
                return "\n\n".join(pages)
            except Exception as e:
                return f"[No se pudo extraer el texto del PDF: {e}]"

        elif file_type == 'docx':
            try:
                doc = Document(io.BytesIO(file_bytes))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                return "\n\n".join(paragraphs)
            except Exception as e:
                return f"[No se pudo extraer el texto del DOCX: {e}]"

    except Exception as e:
        return f"[Error al procesar el archivo: {e}]"

    return "[Formato no soportado]"

def download_youtube_media(url: str, mode: str = 'video', progress_callback=None) -> str:
    """Descarga video o audio de YouTube."""
    logging.info(f"tools.py: Iniciando descarga de {url} (modo: {mode})")
    if not os.path.exists('downloads'):
        os.makedirs('downloads')

    # Iniciar progreso en 0%
    if progress_callback:
        progress_callback(0, 'Iniciando descarga...')

    last_percent = -1
    def hook(d):
        nonlocal last_percent
        if d['status'] == 'downloading':
            percent = 0
            # Método 1: Calcular desde bytes (más fiable)
            downloaded = d.get('downloaded_bytes', 0)
            total = d.get('total_bytes') or d.get('total_bytes_estimate', 0)
            if downloaded and total and total > 0:
                percent = (downloaded / total) * 100
            else:
                # Método 2: Parsear _percent_str limpiando ANSI codes
                try:
                    import re as _re
                    p_str = d.get('_percent_str', '0%')
                    # Limpiar códigos ANSI: \x1b[...m
                    p_clean = _re.sub(r'\x1b\[[0-9;]*m', '', p_str)
                    p_clean = p_clean.replace('%', '').strip()
                    percent = float(p_clean)
                except (ValueError, TypeError):
                    return
            
            percent = min(percent, 99.9)  # Reservar 100% para 'finished'
            if percent >= last_percent + 1:
                last_percent = percent
                if progress_callback:
                    progress_callback(percent, 'Descargando...')
        elif d['status'] == 'finished':
            if progress_callback:
                progress_callback(100, 'Finalizando...')

    ydl_opts = {
        'progress_hooks': [hook],
        'outtmpl': 'downloads/%(title)s.%(ext)s',
        'quiet': True,
        'no_warnings': True,
        'source_address': '0.0.0.0', # Forzar IPv4 para evitar timeouts de red
        'nocheckcertificate': True,
    }

    if mode == 'audio':
        ydl_opts.update({
            'format': 'bestaudio/best',
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'mp3',
                'preferredquality': '192',
            }],
        })
    else:
        ydl_opts.update({
            'format': 'bestvideo[ext=mp4]+bestaudio[ext=m4a]/best[ext=mp4]/best',
        })

    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=True)
            filename = ydl.prepare_filename(info)
            # Si es audio, la extensión cambia a .mp3 por el postprocessor
            if mode == 'audio':
                filename = os.path.splitext(filename)[0] + '.mp3'
            return os.path.basename(filename)
    except Exception as e:
        error_msg = str(e)
        if "ffmpeg" in error_msg.lower():
            return "Error: FFmpeg no encontrado. Es necesario para procesar videos de YouTube. Por favor instálalo."
        return f"Error: {error_msg}"

def create_video_from_folder(folder_input: str, progress_callback=None) -> str:
    """Crea un video MP4 desde imágenes + un MP3 dentro de una carpeta."""
    try:
        if progress_callback:
            progress_callback(5, "Preparando creación de video...")

        folder_input = (folder_input or "").strip().strip('"').strip("'")
        if not folder_input:
            return "Error: No se indicó una carpeta."

        # Resolver ruta: absoluta o relativa al Escritorio del usuario
        if os.path.isabs(folder_input):
            folder_path = Path(folder_input)
        else:
            desktop = Path.home() / "Desktop"
            folder_path = desktop / folder_input

        if not folder_path.exists() or not folder_path.is_dir():
            return f"Error: La carpeta no existe: {folder_path}"

        if progress_callback:
            progress_callback(15, "Buscando imágenes y audio...")

        image_exts = {".jpg", ".jpeg", ".png", ".webp", ".bmp"}
        images = sorted([p for p in folder_path.iterdir() if p.is_file() and p.suffix.lower() in image_exts])
        mp3_files = sorted([p for p in folder_path.iterdir() if p.is_file() and p.suffix.lower() == ".mp3"])

        if not images:
            return "Error: No se encontraron imágenes en la carpeta."
        if not mp3_files:
            return "Error: No se encontró ningún archivo MP3 en la carpeta."

        audio_file = mp3_files[0]
        if not os.path.exists("downloads"):
            os.makedirs("downloads")

        safe_base = re.sub(r"[^\w\-. ]", "", folder_path.name).strip().replace(" ", "_") or "video"
        output_name = f"{safe_base}.mp4"
        output_path = Path("downloads") / output_name

        # Archivo de lista para ffmpeg (concat demuxer)
        list_file = Path("downloads") / f"ffmpeg_list_{uuid.uuid4().hex[:8]}.txt"

        if progress_callback:
            progress_callback(35, "Preparando secuencia de imágenes...")

        with open(list_file, "w", encoding="utf-8") as f:
            for img in images:
                img_posix = img.resolve().as_posix()
                f.write(f"file '{img_posix}'\n")
                f.write("duration 4\n")
            # Repetir última imagen para que ffmpeg respete su duración
            f.write(f"file '{images[-1].resolve().as_posix()}'\n")

        # Duración estimada del video (4s por imagen)
        video_duration = float(len(images) * 4)

        # Intentar leer duración real del audio para calcular fade out suave al final
        audio_duration = 0.0
        try:
            probe_cmd = [
                "ffprobe",
                "-v", "error",
                "-show_entries", "format=duration",
                "-of", "default=noprint_wrappers=1:nokey=1",
                str(audio_file.resolve())
            ]
            probe_res = subprocess.run(probe_cmd, capture_output=True, text=True, encoding="utf-8", errors="replace")
            if probe_res.returncode == 0 and probe_res.stdout.strip():
                audio_duration = float(probe_res.stdout.strip())
        except Exception:
            audio_duration = 0.0

        output_duration = min(video_duration, audio_duration) if audio_duration > 0 else video_duration
        fade_in_duration = 0.6
        fade_out_duration = 0.8

        # Añadimos fade in y fade out visual al video
        video_filter = (
            f"fps=30,format=yuv420p,"
            f"fade=t=in:st=0:d={fade_in_duration:.2f},"
            f"fade=t=out:st={output_duration - fade_out_duration:.2f}:d={fade_out_duration:.2f}"
        )
        audio_filter = (
            f"afade=t=in:st=0:d={fade_in_duration:.2f},"
            f"afade=t=out:st={output_duration - fade_out_duration:.2f}:d={fade_out_duration:.2f}"
        )

        cmd = [
            "ffmpeg",
            "-y",
            "-f", "concat",
            "-safe", "0",
            "-i", str(list_file.resolve()),
            "-i", str(audio_file.resolve()),
            "-vf", video_filter,
            "-af", audio_filter,
            "-c:v", "libx264",
            "-c:a", "aac",
            "-t", f"{output_duration:.2f}",
            str(output_path.resolve())
        ]

        # Hilo para simular avance suave durante la codificación del video
        import threading
        import time

        stop_progress = False
        current_progress = 40.0

        def update_smooth_progress():
            nonlocal current_progress
            while not stop_progress and current_progress < 95.0:
                time.sleep(0.3)
                if stop_progress:
                    break
                # Incremento decreciente a medida que nos acercamos al 95%
                inc = (95.0 - current_progress) * 0.08
                current_progress += max(inc, 0.3)
                if progress_callback:
                    progress_callback(int(current_progress), "Renderizando video y mezclando audio con FFmpeg...")

        if progress_callback:
            progress_callback(35, "Iniciando renderizado con FFmpeg...")

        progress_thread = threading.Thread(target=update_smooth_progress, daemon=True)
        progress_thread.start()

        try:
            result = subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8", errors="replace")
        finally:
            stop_progress = True
            progress_thread.join(timeout=1.0)

        try:
            if list_file.exists():
                list_file.unlink()
        except Exception:
            pass

        if result.returncode != 0:
            err = (result.stderr or "").strip()
            if "ffmpeg" in err.lower():
                return "Error: FFmpeg no encontrado o no disponible en PATH."
            return f"Error: No se pudo crear el video. {err[:400]}"

        if progress_callback:
            progress_callback(100, "Video creado correctamente.")
        return output_name

    except Exception as e:
        return f"Error: {str(e)}"

def execute_python_code(code: str, prompt: str = None) -> dict:
    """Ejecuta código Python localmente de forma segura capturando la salida.
    
    Returns:
        dict con keys: 'stdout', 'stderr', 'success'
    """
    try:
        logging.info("tools.py: Ejecutando script de python generado por IA.")
        # Limpiar posibles comillas markdown residuales
        if code.startswith("```python"):
            code = code[9:]
        if code.startswith("```"):
            code = code[3:]
        if code.endswith("```"):
            code = code[:-3]
        code = code.strip()

        # Linter de sintaxis local antes de ejecutar
        try:
            import ast
            ast.parse(code)
        except SyntaxError as se:
            return {
                "stdout": "",
                "stderr": f"Error de Sintaxis detectado por Linter local en linea {se.lineno}: {se.msg}\nCodigo problemático: {se.text}",
                "success": False
            }

        # Ejecutar en un subproceso con timeout de 60s
        import os
        env = os.environ.copy()
        env["PYTHONIOENCODING"] = "utf-8"
        
        result = subprocess.run(
            [sys.executable, "-c", code],
            capture_output=True,
            text=True,
            timeout=60,
            encoding='utf-8',
            errors='replace',
            env=env
        )
        
        stdout = result.stdout.strip() if result.stdout else ""
        stderr = result.stderr.strip() if result.stderr else ""
        success = result.returncode == 0
        
        if success and prompt:
            save_successful_script(prompt, code)
            
        return {
            "stdout": stdout,
            "stderr": stderr,
            "success": success
        }
    except subprocess.TimeoutExpired:
        return {
            "stdout": "",
            "stderr": "Error: El script tardó demasiado (más de 60 segundos) y fue interrumpido.",
            "success": False
        }
    except Exception as e:
        return {
            "stdout": "",
            "stderr": f"Error crítico al ejecutar: {e}",
            "success": False
        }

def get_youtube_subtitles_clean(url: str) -> str:
    """Intenta descargar subtítulos o transcripción automática de YouTube usando yt-dlp.
    Devuelve el texto limpio o None si no hay subtítulos o si falla."""
    logging.info(f"tools.py: Intentando obtener subtítulos para {url}")
    # Limpiar url de posibles caracteres residuales
    url = url.strip().rstrip('.')
    ydl_opts = {
        'skip_download': True,
        'writesubtitles': True,
        'writeautomaticsub': True,
        'subtitleslangs': ['es', 'es-419', 'en'],
        'quiet': True,
        'no_warnings': True,
        'outtmpl': 'downloads/%(id)s.%(ext)s',
    }
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        try:
            info = ydl.extract_info(url, download=True)
            video_id = info.get('id')
            if not video_id:
                return None
            
            import glob
            # Buscar archivos VTT en downloads que coincidan con el id del video
            files = glob.glob(os.path.join('downloads', f"{video_id}.*.vtt"))
            if not files:
                return None
            
            # Leer el primer archivo encontrado
            vtt_file = files[0]
            logging.info(f"tools.py: Leyendo archivo de subtítulos: {vtt_file}")
            with open(vtt_file, 'r', encoding='utf-8', errors='replace') as f:
                lines = f.readlines()
            
            # Borrar archivos temporales descargados
            for f_path in files:
                try:
                    os.remove(f_path)
                except Exception as ex:
                    logging.error(f"Error al borrar archivo de subtítulos temporal {f_path}: {ex}")
            
            # Limpiar contenido VTT
            text_lines = []
            for line in lines:
                line = line.strip()
                # Omitir metadatos de VTT y marcas de tiempo
                if not line or line.startswith('WEBVTT') or line.startswith('Kind:') or line.startswith('Language:') or '-->' in line or line.isdigit():
                    continue
                # Limpiar las marcas de alineación HTML (como <c>...) y duplicados simples
                clean_line = re.sub(r'<[^>]+>', '', line).strip()
                if clean_line and (not text_lines or text_lines[-1] != clean_line):
                    text_lines.append(clean_line)
                    
            text = " ".join(text_lines)
            # Limitar tamaño de subtítulos para no saturar el contexto de Ollama
            return text[:20000]
        except Exception as e:
            logging.error(f"Error obteniendo subtítulos con yt-dlp: {e}")
            return None

def transcribe_youtube_audio_with_whisper(url: str, progress_callback=None) -> str:
    """Descarga el audio de un video de YouTube y lo transcribe usando Whisper local.
    Devuelve la transcripción limpia o un mensaje de error."""
    logging.info(f"tools.py: Iniciando transcripción con Whisper para {url}")
    
    # 1. Descargar audio con progress_callback
    if progress_callback:
        progress_callback(10, "Descargando audio de YouTube para transcribir...")
    
    # Llamamos a download_youtube_media con mode='audio'
    filename = download_youtube_media(url, mode='audio', progress_callback=progress_callback)
    if "Error" in filename:
        return f"[Error al descargar el audio del video para transcribir: {filename}]"
    
    audio_path = os.path.join('downloads', filename)
    
    if not os.path.exists(audio_path):
        return f"[Error: El archivo de audio descargado no existe en {audio_path}]"
        
    try:
        # 2. Cargar Whisper y transcribir
        if progress_callback:
            progress_callback(40, "Inicializando Whisper local (modelo base)...")
            
        import whisper
        # Cargamos el modelo "base" multilenguaje. Es un buen equilibrio entre rapidez y precisión.
        model = whisper.load_model("base")
        
        if progress_callback:
            progress_callback(60, "Procesando audio y transcribiendo (esto puede tardar un momento)...")
            
        result = model.transcribe(audio_path)
        transcript = result.get("text", "").strip()
        
        if not transcript:
            return "[Error: Whisper no detectó audio o el texto transcrito está vacío]"
            
        if progress_callback:
            progress_callback(95, "Transcripción completada con éxito.")
            
        return transcript
    except Exception as e:
        logging.error(f"Error transcribiendo con Whisper: {e}")
        return f"[Error durante la transcripción con Whisper: {str(e)}]"
    finally:
        # Borrar el archivo de audio temporal de downloads para liberar espacio
        try:
            if os.path.exists(audio_path):
                os.remove(audio_path)
                logging.info(f"tools.py: Archivo de audio temporal borrado: {audio_path}")
        except Exception as ex:
            logging.error(f"Error borrando archivo de audio temporal {audio_path}: {ex}")

